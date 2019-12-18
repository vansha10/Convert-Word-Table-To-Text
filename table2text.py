from docx import Document

tableDocument = Document('table.docx')

textDocument = Document()

for table in tableDocument.tables:
    for row in table.rows:
        s = ""
        for cell in row.cells:
            s += cell.text + " - "
        textDocument.add_paragraph(s)
        
textDocument.save("text.docx")